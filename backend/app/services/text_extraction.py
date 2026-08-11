"""Extracts plain text out of an uploaded file's bytes.

This is the cheap, local half of parsing — no LLM call, just pulling text out
of a PDF/DOCX/TXT/MD so there's something to hand the structured-extraction
step. Runs synchronously right after upload; the LLM call is what gets
deferred to a background task.
"""

from __future__ import annotations

import io
from dataclasses import dataclass

import docx
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.exceptions import AppError


class TextExtractionError(AppError):
    status_code = 422
    code = "text_extraction_failed"


@dataclass(frozen=True)
class ExtractedText:
    text: str
    page_count: int | None
    word_count: int


def _count_words(text: str) -> int:
    return len(text.split())


def _extract_pdf(content: bytes) -> ExtractedText:
    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError as exc:
        raise TextExtractionError(f"Could not open PDF: {exc}") from exc

    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(page.strip() for page in pages if page.strip())

    if not text.strip():
        raise TextExtractionError(
            "No extractable text found in this PDF — it may be a scanned image "
            "without OCR."
        )

    return ExtractedText(text=text, page_count=len(reader.pages), word_count=_count_words(text))


def _extract_docx(content: bytes) -> ExtractedText:
    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001 - python-docx raises plain Exception/KeyError
        raise TextExtractionError(f"Could not open DOCX: {exc}") from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Tables carry real content on many resumes (skills grids, contact blocks).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    if not text.strip():
        raise TextExtractionError("No extractable text found in this DOCX.")

    return ExtractedText(text=text, page_count=None, word_count=_count_words(text))


def _extract_plain_text(content: bytes) -> ExtractedText:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")

    if not text.strip():
        raise TextExtractionError("File is empty.")

    return ExtractedText(text=text, page_count=None, word_count=_count_words(text))


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_plain_text,
    ".md": _extract_plain_text,
}


def extract_text(content: bytes, extension: str) -> ExtractedText:
    """Dispatch to the right extractor for `extension` (e.g. ``.pdf``).

    Raises TextExtractionError for anything unreadable — callers treat that as
    a parse failure (`parse_status = failed`), not a crash.
    """
    extractor = _EXTRACTORS.get(extension.lower())
    if extractor is None:
        raise TextExtractionError(f"No text extractor registered for '{extension}'.")
    return extractor(content)
